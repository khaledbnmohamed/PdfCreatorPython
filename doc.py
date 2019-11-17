from docx import Document
from docx.shared import Inches
import random
from docx.shared import Length
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

class docGenerator:

    def build_pdf(self,name,true_questions,choose_questions,choices_dict):

        document = Document()
        style = document.styles['List Number']
        font = style.font
        font.size = Pt(10)

        # document.add_picture('bahr_logo.png', width=Inches(1.25))

        document.add_heading('Midterm November 2019', 0)
        # p = document.add_paragraph('A plain paragraph having some ')
        # p.add_run('bold').bold = True
        # p.add_run(' and some ')
        # p.add_run('italic.').italic = True

        document.add_heading('Choose only one answer for each of the following question:', level=1)

        random.shuffle(choose_questions)
        choose_questions = choose_questions[0 : 15] 
        for x in choose_questions:
            paragraph = document.add_paragraph(x, style='List Number')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after= Pt(4)
            paragraph_format.keep_with_next= True

            paragraph = document.add_paragraph(choices_dict[x])


        document.save('Choice_model_'+str(name)+'.docx')