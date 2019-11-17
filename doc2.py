from docx import Document
from docx.shared import Inches
import random
from docx.shared import Length
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

class docGenerator:

    def build_pdf(self,name,true_questions,choose_questions,choices_dict):

        document = Document()
        style = document.styles['List Number']
        font = style.font
        font.size = Pt(10)

        # document.add_picture('bahr_logo.png', width=Inches(1.25))

        document.add_heading('Midterm November 2019', 0)
        # p = document.add_paragraph('A plain paragraph having some ')
        # p.add_run('bold').bold = True
        # p.add_run(' and some ')
        # p.add_run('italic.').italic = True


        document.add_heading('Some of the following statements are true and others are false:', level=1)

        random.shuffle(true_questions)
        true_questions = true_questions[0 : 15] 

        for x in true_questions:
            document.add_paragraph(x, style='List Number')

        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        # table = document.add_table(rows=1, cols=3)
        # hdr_cells = table.rows[0].cells
        # hdr_cells[0].text = 'Qty'
        # hdr_cells[1].text = 'Id'
        # hdr_cells[2].text = 'Desc'
        # for qty, id, desc in records:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = str(qty)
        #     row_cells[1].text = id
        #     row_cells[2].text = desc

        document.add_page_break()

        document.save('True_false_model'+str(name)+'.docx')